# -*- coding: utf-8 -*-
"""Parse 生字語詞總複習表.docx / 成語總複習表.docx for each week into
flashcard data, and merge into chinese/data.js (adds vocabCards / idiomCards
to each week object produced by parse_quizzes.py)."""
import json, os, re
from docx import Document

SRC = r"C:\Users\herrs\OneDrive\Documents\中文課\四年級\暑假複習"
GROUPS = ["G1_四上L01-05","G2_四上L06-10","G3_四上L11-14","G4_四下L01-05","G5_四下L06-10","G6_四下L11-14"]
LESSON_RANGE = {1:(1,5), 2:(6,10), 3:(11,14), 4:(1,5), 5:(6,10), 6:(11,14)}
ARABIC = "一二三四五六七八九十十一十二十三十四"

def lesson_label(n):
    return f"第{ARABIC[n-1]}課" if n <= len(ARABIC) else f"第{n}課"

def col_index(header, *names):
    for name in names:
        for i, h in enumerate(header):
            if name in h:
                return i
    return None

def norm_lesson(s):
    s = (s or "").strip()
    if not s:
        return s
    if s.startswith("第") and s.endswith("課"):
        return s
    return f"第{s}課"

def parse_vocab(path, lesson_start, lesson_end):
    doc = Document(path)
    cards = []
    qualifying_tables = []
    for t in doc.tables:
        header = [c.text.strip() for c in t.rows[0].cells]
        if any("語詞" in h for h in header) and any("簡單英文解釋" in h for h in header):
            qualifying_tables.append((t, header))

    n_lessons = lesson_end - lesson_start + 1
    infer_lesson_by_table = (len(qualifying_tables) == n_lessons) and \
        not any(col_index(h, "課別") is not None for _, h in qualifying_tables)

    for ti, (t, header) in enumerate(qualifying_tables):
        i_lesson = col_index(header, "課別")
        i_char = col_index(header, "生字")
        i_term = col_index(header, "語詞")
        i_zhuyin = col_index(header, "注音")
        i_en = col_index(header, "簡單英文解釋")
        i_ex = col_index(header, "例句")
        fallback_lesson = lesson_label(lesson_start + ti) if infer_lesson_by_table else None
        last_lesson, last_char = "", ""
        for row in t.rows[1:]:
            cells = [c.text.strip() for c in row.cells]
            term = cells[i_term] if i_term is not None else ""
            if not term:
                continue
            lesson_val = (cells[i_lesson] if i_lesson is not None else fallback_lesson) or ""
            char_val = cells[i_char] if i_char is not None else ""
            if lesson_val:
                last_lesson = lesson_val
            else:
                lesson_val = last_lesson
            if char_val:
                last_char = char_val
            else:
                char_val = last_char
            cards.append({
                "lesson": norm_lesson(lesson_val),
                "char": char_val,
                "term": term,
                "zhuyin": cells[i_zhuyin] if i_zhuyin is not None else "",
                "en": cells[i_en] if i_en is not None else "",
                "example": cells[i_ex] if i_ex is not None else "",
            })
    return cards

def parse_idiom(path):
    doc = Document(path)
    cards = []
    for t in doc.tables:
        header = [c.text.strip() for c in t.rows[0].cells]
        if any("成語" in h for h in header) and any("意思" in h for h in header) and any("例句" in h for h in header):
            i_lesson = col_index(header, "課別")
            i_term = col_index(header, "成語")
            i_mean = col_index(header, "意思")
            i_ex = col_index(header, "例句")
            last_lesson = ""
            for row in t.rows[1:]:
                cells = [c.text.strip() for c in row.cells]
                term = cells[i_term] if i_term is not None else ""
                if not term:
                    continue
                lesson_val = (cells[i_lesson] if i_lesson is not None else "") or ""
                if lesson_val:
                    last_lesson = lesson_val
                else:
                    lesson_val = last_lesson
                cards.append({
                    "lesson": norm_lesson(lesson_val),
                    "term": term,
                    "meaning": cells[i_mean] if i_mean is not None else "",
                    "example": cells[i_ex] if i_ex is not None else "",
                })
    return cards

weeks_cards = {}
for idx, g in enumerate(GROUPS, start=1):
    ls, le = LESSON_RANGE[idx]
    vocab_path = os.path.join(SRC, g, "生字語詞總複習表.docx")
    idiom_path = os.path.join(SRC, g, "成語總複習表.docx")
    weeks_cards[idx] = {
        "vocabCards": parse_vocab(vocab_path, ls, le),
        "idiomCards": parse_idiom(idiom_path),
    }

# merge into existing chinese/data.js (which defines CHINESE_WEEKS as a JSON-ish array)
data_js_path = r"C:\Users\herrs\Claude\valerie-summer-review\chinese\data.js"
with open(data_js_path, encoding="utf-8") as f:
    content = f.read()
m = re.search(r"const CHINESE_WEEKS = (\[.*\]);", content, re.DOTALL)
weeks = json.loads(m.group(1))
for w in weeks:
    extra = weeks_cards[w["id"]]
    w["vocabCards"] = extra["vocabCards"]
    w["idiomCards"] = extra["idiomCards"]

with open(data_js_path, "w", encoding="utf-8") as f:
    f.write("const CHINESE_WEEKS = ")
    f.write(json.dumps(weeks, ensure_ascii=False, indent=2))
    f.write(";\n")

with open(r"C:\Users\herrs\Claude\valerie-summer-review\_cards_log.txt", "w", encoding="utf-8") as f:
    for idx in range(1, 7):
        f.write(f"week{idx}: vocabCards={len(weeks_cards[idx]['vocabCards'])} idiomCards={len(weeks_cards[idx]['idiomCards'])}\n")
