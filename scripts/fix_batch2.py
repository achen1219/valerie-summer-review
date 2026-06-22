# -*- coding: utf-8 -*-
"""Batch fix for user-reported zhuyin/pairing errors (拆/折, 劃, 量, 柴, 嚼,
積, 調, 螺/累/壘, 圈/卷/捲, 匙, 趙/擔), plus user-supplied corrections for 宿, 調."""
from docx import Document

PATH = r"C:\Users\herrs\OneDrive\Documents\中文課\四年級\暑假複習\字音字形總複習.docx"
doc = Document(PATH)
t1 = doc.tables[0]  # 多音字: 字 | 課別 | 讀音與例詞
t2 = doc.tables[1]  # 形近字: 課別 | 字組 | 字／注音／例句

DUOYIN_FIXES = {
    "宿": "ㄙㄨˋ／過夜居住、舊的積久的，例：借宿、住宿、宿怨、宿疾\nㄒㄧㄡˋ／星座，例：星宿\nㄒㄧㄡˇ／夜晚，例：一宿",
    "調": "ㄊㄧㄠˊ／使和諧均勻、和解、混合配合、嘲笑挑逗，例：調和、調解、調色、調侃\nㄉㄧㄠˋ／職務更動、派遣、互換、提取、樂律、語言聲調，例：調職、調兵遣將、對調、音調",
    "劃": "ㄏㄨㄚˋ／分開、分界、設計籌謀，例：劃分、籌劃、計劃\nㄏㄨㄚˊ／用利器拖拉過、擦掠，例：劃了一道傷口、劃火柴",
    "量": "ㄌㄧㄤˊ／用工具測量，例：量身高、量杯\nㄌㄧㄤˋ／估計、衡量、數量，例：量入為出、不自量力、含量、重量、容量、度量衡",
    "柴": "ㄔㄞˊ／供燃燒用的小木枯枝，例：打柴、薪柴\nㄓㄞˋ／柵欄，通「寨」，例：鹿柴",
    "嚼": "ㄐㄩㄝˊ／咀嚼、嚼蠟\nㄐㄧㄠˊ／窮嚼、細嚼慢嚥",
}
DELETE_DUOYIN_CHARS = {"折"}  # not a real multi-tone reading; was actually 拆 mislabeled

for row in t1.rows[1:]:
    char = row.cells[0].text.strip()
    if char in DUOYIN_FIXES:
        row.cells[2].text = DUOYIN_FIXES[char]

# delete the wrong 折 row from table 1
for row in list(t1.rows[1:]):
    if row.cells[0].text.strip() in DELETE_DUOYIN_CHARS:
        row._element.getparent().remove(row._element)

XINGJIN_LABEL_FIXES = {
    "積／績／蹟": ["積：ㄐㄧ／堆積、日積月累", "績：ㄐㄧˋ／成績、豐功偉績", "蹟：ㄐㄧˋ／奇蹟、名勝古蹟"],
    "其他單字（趙／擔）": ["趙：ㄓㄠˋ／完璧歸趙", "擔：ㄉㄢ／擔心、擔任、承擔", "擔：ㄉㄢˋ／擔子、重擔"],
    "調／凋／周": ["調：ㄊㄧㄠˊ／調和、調色、調解", "調：ㄉㄧㄠˋ／調職、調換、對調", "凋：ㄉㄧㄠ／凋萎、凋零", "周：ㄓㄡ／四周、周圍"],
    "螺／累／壘": ["螺：ㄌㄨㄛˊ／陀螺、螺絲", "累：ㄌㄟˋ／疲累、勞累", "累：ㄌㄟˇ／累積、經年累月", "壘：ㄌㄟˇ／壘球、堡壘"],
    "其他補充字（匙）": ["匙：ㄔˊ／湯匙、茶匙", "匙：˙ㄕ／鑰匙"],
}

for row in t2.rows[1:]:
    label = row.cells[1].text.strip()
    if label in XINGJIN_LABEL_FIXES:
        row.cells[2].text = "\n".join(XINGJIN_LABEL_FIXES[label])
    elif label == "圈／卷／捲":
        lines = row.cells[2].text.split("\n")
        lines.insert(0, "圈：ㄐㄩㄢˋ／牛圈、羊圈")
        lines.insert(0, "圈：ㄑㄩㄢ／圓圈、圈起來")
        row.cells[2].text = "\n".join(lines)

# add new 形近字 group: 拆／折 (under 四上L6, replacing the deleted wrong 多音字 entry)
new_row = t2.add_row()
new_row.cells[0].text = "四上L6"
new_row.cells[1].text = "拆／折"
new_row.cells[2].text = "拆：ㄔㄞ／拆除、拆信、拆閱\n折：ㄓㄜˊ／骨折、一波三折"

doc.save(PATH)
print("done")
