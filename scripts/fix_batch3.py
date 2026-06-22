# -*- coding: utf-8 -*-
"""Move 匙 from 形近字 (其他補充字) into 多音字, where it actually belongs
(it's one character with two readings, not a confusable-character pair)."""
from docx import Document

PATH = r"C:\Users\herrs\OneDrive\Documents\中文課\四年級\暑假複習\字音字形總複習.docx"
doc = Document(PATH)
t1 = doc.tables[0]  # 多音字: 字 | 課別 | 讀音與例詞
t2 = doc.tables[1]  # 形近字: 課別 | 字組 | 字／注音／例句

# remove the 匙 row from 形近字
for row in list(t2.rows[1:]):
    if row.cells[1].text.strip() == "其他補充字（匙）":
        row._element.getparent().remove(row._element)

# add 匙 to 多音字
new_row = t1.add_row()
new_row.cells[0].text = "匙"
new_row.cells[1].text = "四下L10"
new_row.cells[2].text = "ㄔˊ／湯匙、茶匙\n˙ㄕ／鑰匙"

doc.save(PATH)
print("done")
