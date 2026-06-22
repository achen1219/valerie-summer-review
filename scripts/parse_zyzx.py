# -*- coding: utf-8 -*-
"""Parse 字音字形總複習.docx back into JSON for the website (avoids retyping)."""
import json
from docx import Document

SRC = r"C:\Users\herrs\OneDrive\Documents\中文課\四年級\暑假複習\字音字形總複習.docx"
doc = Document(SRC)

duoyin = []
xingjin = []

t1 = doc.tables[0]  # 字 | 課別 | 讀音與例詞
for row in t1.rows[1:]:
    cells = [c.text for c in row.cells]
    char, lesson, readings_text = cells[0], cells[1], cells[2]
    readings = [r for r in readings_text.split("\n") if r.strip()]
    duoyin.append({"char": char.strip(), "lesson": lesson.strip(), "readings": readings})

t2 = doc.tables[1]  # 課別 | 字組 | 字／注音／例句
for row in t2.rows[1:]:
    cells = [c.text for c in row.cells]
    lesson, label, items_text = cells[0], cells[1], cells[2]
    items = [it for it in items_text.split("\n") if it.strip()]
    xingjin.append({"lesson": lesson.strip(), "label": label.strip(), "items": items})

out_path = r"C:\Users\herrs\Claude\valerie-summer-review\chinese\zyzx-data.js"
with open(out_path, "w", encoding="utf-8") as f:
    f.write("const ZYZX_DUOYIN = ")
    f.write(json.dumps(duoyin, ensure_ascii=False, indent=2))
    f.write(";\nconst ZYZX_XINGJIN = ")
    f.write(json.dumps(xingjin, ensure_ascii=False, indent=2))
    f.write(";\n")

print(f"duoyin={len(duoyin)} xingjin={len(xingjin)}")
